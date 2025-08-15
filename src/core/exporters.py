"""
Tutorial Export Functionality
Generates shareable documents from captured tutorial data
"""

import os
import json
from pathlib import Path
from typing import List, Optional, Dict, Any
from datetime import datetime
import base64
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from PIL import Image

from .storage import TutorialMetadata, TutorialStep, TutorialStorage
from .screenshot_processor import ClickHighlighter
from ..utils.file_utils import sanitize_filename, format_duration


class HTMLExporter:
    """Export tutorial to HTML format with editing capabilities"""
    
    def __init__(self):
        self.template = self._get_html_template()
        self.click_highlighter = ClickHighlighter()
    
    def export(self, metadata: TutorialMetadata, steps: List[TutorialStep], 
               project_path: Path) -> str:
        """
        Export tutorial to HTML format
        
        Args:
            metadata: Tutorial metadata
            steps: List of tutorial steps
            project_path: Path to tutorial project directory
            
        Returns:
            Path to generated HTML file
        """
        # Use sanitized tutorial title as filename
        safe_title = sanitize_filename(metadata.title or "untitled")
        output_path = project_path / "output" / f"{safe_title}.html"
        
        # Convert steps to HTML
        steps_html = self._generate_steps_html(steps, project_path)
        
        # Generate complete HTML
        click_css = self.click_highlighter.get_click_indicator_css()
        html_content = self.template.format(
            title=metadata.title,
            description=metadata.description,
            created_date=datetime.fromtimestamp(metadata.created_at).strftime("%B %d, %Y"),
            step_count=metadata.step_count,
            duration=format_duration(metadata.duration),
            steps_html=steps_html,
            click_css=click_css
        )
        
        # Write HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return str(output_path)
    
    def _generate_steps_html(self, steps: List[TutorialStep], project_path: Path) -> str:
        """Generate HTML for tutorial steps"""
        steps_html = ""
        
        for step in steps:
            # Convert screenshot to base64 for embedding
            screenshot_b64 = ""
            click_indicator_html = ""
            screenshot_width = 0
            screenshot_height = 0
            
            if step.screenshot_path:
                screenshot_full_path = project_path / step.screenshot_path
                if screenshot_full_path.exists():
                    # Load image to get dimensions
                    with Image.open(screenshot_full_path) as img:
                        screenshot_width, screenshot_height = img.size
                    
                    with open(screenshot_full_path, 'rb') as img_file:
                        screenshot_b64 = base64.b64encode(img_file.read()).decode('utf-8')
                    
                    # Generate click indicator using percentage coordinates if available
                    if step.coordinates_pct and screenshot_width > 0:
                        # Use percentage coordinates for accurate positioning
                        pixel_x = int(step.coordinates_pct[0] * screenshot_width)
                        pixel_y = int(step.coordinates_pct[1] * screenshot_height)
                        click_indicator_html = self.click_highlighter.add_animated_click_indicator_html(
                            pixel_x, pixel_y, 
                            screenshot_width, screenshot_height
                        )
                    elif step.coordinates and screenshot_width > 0:
                        # Fallback to absolute coordinates for legacy data
                        click_indicator_html = self.click_highlighter.add_animated_click_indicator_html(
                            step.coordinates[0], step.coordinates[1], 
                            screenshot_width, screenshot_height
                        )
            
            # Generate step HTML with click indicators
            screenshot_html = ""
            if screenshot_b64:
                screenshot_html = f"""
                <div class="screenshot-container">
                    <img src="data:image/png;base64,{screenshot_b64}" 
                         alt="Step {step.step_number} screenshot" 
                         class="step-screenshot">
                    {click_indicator_html}
                </div>
                """
            
            step_html = f"""
            <div class="tutorial-step" data-step-id="{step.step_id}">
                <div class="step-header">
                    <span class="step-number">{step.step_number}</span>
                    <button class="delete-step" onclick="deleteStep('{step.step_id}')">×</button>
                </div>
                <div class="step-content">
                    <div class="step-description" contenteditable="true">{step.description}</div>
                    {screenshot_html}
                    <div class="step-metadata">
                        {f'<span class="coordinates">Clicked at: ({int(step.coordinates[0])}, {int(step.coordinates[1])})</span>' if step.coordinates else ''}
                        {f'<span class="ocr-confidence">OCR Confidence: {step.ocr_confidence:.1%}</span>' if step.ocr_confidence > 0 else ''}
                    </div>
                </div>
            </div>
            """
            steps_html += step_html
        
        return steps_html
    
    
    def _get_html_template(self) -> str:
        """Get the HTML template for tutorials"""
        return """<!DOCTYPE html>
<!-- This is a standalone exported tutorial file -->
<!-- Export buttons are disabled in standalone mode -->
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} - Tutorial</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        
        {click_css}
        
        .tutorial-header {{
            border-bottom: 2px solid #e1e5e9;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        
        .tutorial-title {{
            font-size: 2.5em;
            margin: 0;
            color: #2c3e50;
        }}
        
        .tutorial-meta {{
            margin-top: 10px;
            color: #666;
            font-size: 0.9em;
        }}
        
        .tutorial-step {{
            background: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 8px;
            margin-bottom: 20px;
            overflow: hidden;
            position: relative;
        }}
        
        .step-header {{
            background: #007bff;
            color: white;
            padding: 10px 15px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        
        .step-number {{
            font-weight: bold;
            font-size: 1.1em;
        }}
        
        .delete-step {{
            background: rgba(255,255,255,0.2);
            border: none;
            color: white;
            width: 25px;
            height: 25px;
            border-radius: 50%;
            cursor: pointer;
            font-size: 16px;
            display: flex;
            align-items: center;
            justify-content: center;
        }}
        
        .delete-step:hover {{
            background: rgba(255,255,255,0.3);
        }}
        
        .step-content {{
            padding: 20px;
        }}
        
        .step-description {{
            font-size: 1.1em;
            margin-bottom: 15px;
            min-height: 25px;
            padding: 5px;
            border-radius: 4px;
        }}
        
        .step-description:focus {{
            outline: 2px solid #007bff;
            background: white;
        }}
        
        .step-screenshot {{
            max-width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 4px;
            margin-bottom: 10px;
        }}
        
        .step-metadata {{
            font-size: 0.85em;
            color: #666;
            display: flex;
            gap: 15px;
        }}
        
        .export-controls {{
            margin-top: 30px;
            padding: 20px;
            background: #f1f3f5;
            border-radius: 8px;
            text-align: center;
        }}
        
        .export-btn {{
            background: #28a745;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 5px;
            margin: 0 5px;
            cursor: pointer;
            font-size: 0.9em;
        }}
        
        .export-btn:hover {{
            background: #218838;
        }}
        
        .deleted {{
            opacity: 0.5;
            transition: opacity 0.3s;
        }}
    </style>
</head>
<body>
    <div class="tutorial-header">
        <h1 class="tutorial-title">{title}</h1>
        <div class="tutorial-meta">
            <p>{description}</p>
            <p><strong>Created:</strong> {created_date} | <strong>Steps:</strong> {step_count} | <strong>Duration:</strong> {duration}</p>
        </div>
    </div>
    
    <div class="tutorial-steps">
        {steps_html}
    </div>
    
    <div class="export-controls">
        <h3>Standalone Tutorial</h3>
        <p style="color: #666; margin: 0;">This is an exported tutorial file. To edit or re-export, use the TutorialMaker web interface.</p>
        <p style="color: #666; margin: 5px 0 0 0; font-size: 0.9em;">
            <strong>💡 Tip:</strong> To edit this tutorial:<br>
            • <strong>Development:</strong> Run <code style="background: #f0f0f0; padding: 2px 4px; border-radius: 3px;">python main.py</code><br>
            • <strong>Production:</strong> Run <code style="background: #f0f0f0; padding: 2px 4px; border-radius: 3px;">./tutorialmaker</code> executable
        </p>
    </div>

    <script>
        // Standalone mode - disable all editing functionality
        document.addEventListener('DOMContentLoaded', function() {{
            // Remove delete buttons
            document.querySelectorAll('.delete-step').forEach(btn => {{
                btn.style.display = 'none';
            }});
            
            // Make descriptions read-only
            document.querySelectorAll('.step-description').forEach(desc => {{
                desc.setAttribute('contenteditable', 'false');
                desc.style.cursor = 'default';
                desc.style.backgroundColor = 'transparent';
            }});
            
            // Add notice for users who try to click
            document.addEventListener('click', function(e) {{
                if (e.target.classList.contains('step-description') || e.target.closest('.delete-step')) {{
                    e.preventDefault();
                    if (confirm('This is a standalone exported file. Would you like instructions on how to edit tutorials?')) {{
                        alert('To edit tutorials:\\n\\nDevelopment:\\n1. Run: python main.py\\n2. Browser opens automatically\\n\\nProduction:\\n1. Run: ./tutorialmaker executable\\n2. Browser opens automatically');
                    }}
                }}
            }});
        }});
        
        // Disabled functions (kept for compatibility but non-functional)
        function deleteStep(stepId) {{
            // Disabled in standalone mode
        }}
        
        function exportTutorial(format) {{
            // Disabled in standalone mode
        }}
        
        function saveTutorial() {{
            // Disabled in standalone mode
        }}
        
        // Intersection Observer for click animations
        document.addEventListener('DOMContentLoaded', function() {{
            const observer = new IntersectionObserver((entries) => {{
                entries.forEach(entry => {{
                    if (entry.isIntersecting) {{
                        entry.target.classList.add('step-visible');
                        // Trigger animation with slight delay for better effect
                        setTimeout(() => {{
                            const clickCircles = entry.target.querySelectorAll('.click-circle');
                            clickCircles.forEach(circle => {{
                                // Reset animation by removing and re-adding the class
                                circle.style.animation = 'none';
                                circle.offsetHeight; // Force reflow
                                circle.style.animation = 'clickContract 1.5s ease-out';
                            }});
                        }}, 200);
                    }} else {{
                        entry.target.classList.remove('step-visible');
                    }}
                }});
            }}, {{
                threshold: 0.3, // Trigger when 30% of step is visible
                rootMargin: '0px 0px -50px 0px' // Account for header/footer
            }});
            
            // Observe all tutorial steps
            document.querySelectorAll('.tutorial-step').forEach(step => {{
                observer.observe(step);
            }});
        }});
    </script>
</body>
</html>"""


class WordExporter:
    """Export tutorial to Word document format"""
    
    def __init__(self):
        self.click_highlighter = ClickHighlighter()
    
    def export(self, metadata: TutorialMetadata, steps: List[TutorialStep], 
               project_path: Path) -> str:
        """
        Export tutorial to Word document
        
        Args:
            metadata: Tutorial metadata
            steps: List of tutorial steps
            project_path: Path to tutorial project directory
            
        Returns:
            Path to generated Word document
        """
        # Use sanitized tutorial title as filename
        safe_title = sanitize_filename(metadata.title or "untitled")
        output_path = project_path / "output" / f"{safe_title}.docx"
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(metadata.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Created: {datetime.fromtimestamp(metadata.created_at).strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Steps: {metadata.step_count}")
        doc.add_paragraph(f"Duration: {format_duration(metadata.duration)}")
        
        if metadata.description:
            doc.add_paragraph(metadata.description)
        
        doc.add_page_break()
        
        # Add steps
        for step in steps:
            # Step heading
            doc.add_heading(f"Step {step.step_number}", level=1)
            
            # Step description
            doc.add_paragraph(step.description)
            
            # Add screenshot if available
            if step.screenshot_path:
                screenshot_full_path = project_path / step.screenshot_path
                if screenshot_full_path.exists():
                    try:
                        # Load and process image with click highlighting
                        with Image.open(screenshot_full_path) as img:
                            if step.coordinates_pct:
                                # Use percentage coordinates for accurate positioning
                                img_width, img_height = img.size
                                pixel_x = int(step.coordinates_pct[0] * img_width)
                                pixel_y = int(step.coordinates_pct[1] * img_height)
                                
                                # Add click indicator to image
                                img_with_click = self.click_highlighter.add_click_indicator(
                                    img, pixel_x, pixel_y
                                )
                                
                                # Save processed image temporarily
                                temp_path = project_path / "temp" / f"highlighted_{step.step_id}.png"
                                temp_path.parent.mkdir(exist_ok=True)
                                img_with_click.save(temp_path)
                                
                                # Add processed image to document
                                doc.add_picture(str(temp_path), width=Inches(6))
                                
                                # Clean up temp file
                                temp_path.unlink(missing_ok=True)
                            elif step.coordinates:
                                # Fallback to absolute coordinates for legacy data
                                img_with_click = self.click_highlighter.add_click_indicator(
                                    img, step.coordinates[0], step.coordinates[1]
                                )
                                
                                # Save processed image temporarily
                                temp_path = project_path / "temp" / f"highlighted_{step.step_id}.png"
                                temp_path.parent.mkdir(exist_ok=True)
                                img_with_click.save(temp_path)
                                
                                # Add processed image to document
                                doc.add_picture(str(temp_path), width=Inches(6))
                                
                                # Clean up temp file
                                temp_path.unlink(missing_ok=True)
                            else:
                                # Add original image if no coordinates
                                doc.add_picture(str(screenshot_full_path), width=Inches(6))
                    except Exception as e:
                        doc.add_paragraph(f"[Screenshot: {step.screenshot_path} - Error loading: {e}]")
            
            # Add coordinates and OCR info
            if step.coordinates:
                doc.add_paragraph(f"Click coordinates: ({int(step.coordinates[0])}, {int(step.coordinates[1])})")
            
            if step.ocr_confidence > 0:
                doc.add_paragraph(f"OCR confidence: {step.ocr_confidence:.1%}")
            
            doc.add_paragraph("")  # Space between steps
        
        # Save document
        doc.save(output_path)
        return str(output_path)
    


class PDFExporter:
    """Export tutorial to PDF format"""
    
    def export(self, metadata: TutorialMetadata, steps: List[TutorialStep], 
               project_path: Path) -> str:
        """
        Export tutorial to PDF document
        
        Args:
            metadata: Tutorial metadata
            steps: List of tutorial steps
            project_path: Path to tutorial project directory
            
        Returns:
            Path to generated PDF document
        """
        # Use sanitized tutorial title as filename
        safe_title = sanitize_filename(metadata.title or "untitled")
        output_path = project_path / "output" / f"{safe_title}.pdf"
        
        # Create PDF
        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter
        
        # Title page
        c.setFont("Helvetica-Bold", 24)
        c.drawString(width/2 - len(metadata.title)*6, height - 100, metadata.title)
        
        c.setFont("Helvetica", 12)
        y = height - 150
        c.drawString(50, y, f"Created: {datetime.fromtimestamp(metadata.created_at).strftime('%B %d, %Y')}")
        y -= 20
        c.drawString(50, y, f"Steps: {metadata.step_count}")
        y -= 20
        c.drawString(50, y, f"Duration: {format_duration(metadata.duration)}")
        
        if metadata.description:
            y -= 40
            c.drawString(50, y, "Description:")
            y -= 20
            c.drawString(50, y, metadata.description)
        
        c.showPage()  # New page for steps
        
        # Add steps
        for step in steps:
            y = height - 50
            
            # Step heading
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y, f"Step {step.step_number}")
            y -= 30
            
            # Step description
            c.setFont("Helvetica", 12)
            c.drawString(50, y, step.description)
            y -= 30
            
            # Add screenshot if available
            if step.screenshot_path:
                screenshot_full_path = project_path / step.screenshot_path
                if screenshot_full_path.exists():
                    try:
                        # Load and resize image
                        img = Image.open(screenshot_full_path)
                        img_width, img_height = img.size
                        
                        # Scale to fit page width (max 500px)
                        max_width = 500
                        if img_width > max_width:
                            scale = max_width / img_width
                            img_width = int(img_width * scale)
                            img_height = int(img_height * scale)
                        
                        # Check if image fits on current page
                        if y - img_height < 100:
                            c.showPage()
                            y = height - 50
                        
                        # Draw image
                        c.drawImage(str(screenshot_full_path), 50, y - img_height, 
                                  width=img_width, height=img_height)
                        y -= img_height + 20
                        
                    except Exception as e:
                        c.drawString(50, y, f"[Screenshot error: {e}]")
                        y -= 20
            
            # Add coordinates and OCR info
            if step.coordinates:
                c.setFont("Helvetica", 10)
                c.drawString(50, y, f"Click coordinates: ({int(step.coordinates[0])}, {int(step.coordinates[1])})")
                y -= 15
            
            if step.ocr_confidence > 0:
                c.drawString(50, y, f"OCR confidence: {step.ocr_confidence:.1%}")
                y -= 15
            
            # Start new page if needed
            if y < 150:
                c.showPage()
        
        c.save()
        return str(output_path)
    


class MarkdownExporter:
    """Export tutorial to Markdown format"""
    
    def export(self, metadata: TutorialMetadata, steps: List[TutorialStep], 
               project_path: Path) -> str:
        """
        Export tutorial to Markdown format
        
        Args:
            metadata: Tutorial metadata
            steps: List of tutorial steps
            project_path: Path to tutorial project directory
            
        Returns:
            Path to generated Markdown file
        """
        # Use sanitized tutorial title as filename
        safe_title = sanitize_filename(metadata.title or "untitled")
        output_path = project_path / "output" / f"{safe_title}.md"
        
        # Generate Markdown content
        md_content = self._generate_markdown_content(metadata, steps, project_path)
        
        # Write Markdown file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return str(output_path)
    
    def _generate_markdown_content(self, metadata: TutorialMetadata, 
                                   steps: List[TutorialStep], project_path: Path) -> str:
        """Generate Markdown content for the tutorial"""
        lines = []
        
        # Title and metadata
        lines.append(f"# {metadata.title}")
        lines.append("")
        
        if metadata.description:
            lines.append(f"**Description:** {metadata.description}")
            lines.append("")
        
        lines.append(f"**Created:** {datetime.fromtimestamp(metadata.created_at).strftime('%B %d, %Y')}")
        lines.append(f"**Steps:** {metadata.step_count}")
        lines.append(f"**Duration:** {format_duration(metadata.duration)}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Table of contents
        if len(steps) > 3:
            lines.append("## Table of Contents")
            lines.append("")
            for i, step in enumerate(steps, 1):
                step_title = step.description[:50] + "..." if len(step.description) > 50 else step.description
                lines.append(f"{i}. [{step_title}](#step-{i})")
            lines.append("")
            lines.append("---")
            lines.append("")
        
        # Steps
        for i, step in enumerate(steps, 1):
            lines.append(f"## Step {i}")
            lines.append("")
            lines.append(f"**Action:** {step.description}")
            lines.append("")
            
            # Add coordinates if available
            if step.coordinates:
                lines.append(f"**Click Position:** ({step.coordinates[0]}, {step.coordinates[1]})")
                lines.append("")
            
            # Add OCR text if available
            if step.ocr_text and step.ocr_text.strip():
                lines.append(f"**Detected Text:** \"{step.ocr_text.strip()}\"")
                if step.ocr_confidence and step.ocr_confidence > 0:
                    lines.append(f"**Confidence:** {step.ocr_confidence * 100:.1f}%")
                lines.append("")
            
            # Add screenshot as image reference
            if step.screenshot_path:
                # Handle both full paths and just filenames
                if "/" in step.screenshot_path:
                    # Full path like "screenshots/step_001.png"
                    relative_screenshot_path = f"../{step.screenshot_path}"
                else:
                    # Just filename like "step_001.png"
                    relative_screenshot_path = f"../screenshots/{step.screenshot_path}"
                
                # Check if file exists
                screenshot_full_path = project_path / "screenshots" / (step.screenshot_path.split("/")[-1] if "/" in step.screenshot_path else step.screenshot_path)
                if screenshot_full_path.exists():
                    lines.append(f"![Step {i} Screenshot]({relative_screenshot_path})")
                    lines.append("")
                    lines.append(f"*Screenshot: {relative_screenshot_path}*")
                    lines.append("")
            
            # Add keyboard input if available
            if hasattr(step, 'keyboard_input') and step.keyboard_input:
                lines.append(f"**Typed:** `{step.keyboard_input}`")
                lines.append("")
            
            lines.append("---")
            lines.append("")
        
        # Footer
        lines.append("## Tutorial Information")
        lines.append("")
        lines.append("This tutorial was generated automatically using TutorialMaker.")
        lines.append(f"- **Total Steps:** {len(steps)}")
        lines.append(f"- **Total Duration:** {format_duration(metadata.duration)}")
        lines.append(f"- **Generated on:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        return "\n".join(lines)
    


class TutorialExporter:
    """Main tutorial exporter that coordinates all export formats"""
    
    def __init__(self, storage: TutorialStorage):
        self.storage = storage
        self.html_exporter = HTMLExporter()
        self.word_exporter = WordExporter()
        self.pdf_exporter = PDFExporter()
        self.markdown_exporter = MarkdownExporter()
    
    def export_tutorial(self, tutorial_id: str, formats: List[str] = None) -> Dict[str, str]:
        """
        Export tutorial to specified formats
        
        Args:
            tutorial_id: Tutorial ID to export
            formats: List of formats to export ('html', 'word', 'pdf', 'markdown'). 
                    If None, exports to HTML and Word by default.
        
        Returns:
            Dictionary mapping format names to output file paths
        """
        if formats is None:
            formats = ['html', 'word']
        
        # Load tutorial data
        metadata = self.storage.load_tutorial_metadata(tutorial_id)
        steps = self.storage.load_tutorial_steps(tutorial_id)
        project_path = self.storage.get_project_path(tutorial_id)
        
        if not metadata or not steps or not project_path:
            raise ValueError(f"Tutorial not found or incomplete: {tutorial_id}")
        
        # Ensure output directory exists
        output_dir = project_path / "output"
        output_dir.mkdir(exist_ok=True)
        
        results = {}
        
        # Export to requested formats
        if 'html' in formats:
            try:
                results['html'] = self.html_exporter.export(metadata, steps, project_path)
                print(f"HTML export completed: {results['html']}")
            except Exception as e:
                print(f"HTML export failed: {e}")
                results['html'] = f"Error: {e}"
        
        if 'word' in formats:
            try:
                results['word'] = self.word_exporter.export(metadata, steps, project_path)
                print(f"Word export completed: {results['word']}")
            except Exception as e:
                print(f"Word export failed: {e}")
                results['word'] = f"Error: {e}"
        
        if 'pdf' in formats:
            try:
                results['pdf'] = self.pdf_exporter.export(metadata, steps, project_path)
                print(f"PDF export completed: {results['pdf']}")
            except Exception as e:
                print(f"PDF export failed: {e}")
                results['pdf'] = f"Error: {e}"
        
        if 'markdown' in formats:
            try:
                results['markdown'] = self.markdown_exporter.export(metadata, steps, project_path)
                print(f"Markdown export completed: {results['markdown']}")
            except Exception as e:
                print(f"Markdown export failed: {e}")
                results['markdown'] = f"Error: {e}"
        
        return results
    
    def export_all_tutorials(self, formats: List[str] = None, max_workers: int = 3) -> Dict[str, Dict[str, str]]:
        """
        Export all tutorials to specified formats using concurrent processing
        
        Args:
            formats: List of formats to export
            max_workers: Maximum number of concurrent export operations
            
        Returns:
            Dictionary mapping tutorial IDs to export results
        """
        tutorials = self.storage.list_tutorials()
        results = {}
        
        if not tutorials:
            return results
        
        def export_single_tutorial(tutorial):
            """Export a single tutorial - used by thread pool"""
            try:
                return tutorial.tutorial_id, self.export_tutorial(tutorial.tutorial_id, formats)
            except Exception as e:
                print(f"Failed to export tutorial {tutorial.tutorial_id}: {e}")
                return tutorial.tutorial_id, {"error": str(e)}
        
        # Use ThreadPoolExecutor for concurrent exports
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all export tasks
            future_to_tutorial = {
                executor.submit(export_single_tutorial, tutorial): tutorial 
                for tutorial in tutorials
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_tutorial):
                tutorial = future_to_tutorial[future]
                try:
                    tutorial_id, export_result = future.result()
                    results[tutorial_id] = export_result
                    print(f"Completed export for tutorial: {tutorial.title} ({tutorial_id})")
                except Exception as e:
                    print(f"Export task failed for tutorial {tutorial.tutorial_id}: {e}")
                    results[tutorial.tutorial_id] = {"error": str(e)}
        
        return results